import streamlit as st
import pandas as pd
import gspread
from google.oauth2.service_account import Credentials
import docx
import os
import re
import io
import fitz  # PyMuPDF para converter PDF em imagem

# Configuração da página
st.set_page_config(page_title="Painel de Gestão - Defesa do Idoso SP", layout="wide")

# --- CONEXÃO COM O GOOGLE SHEETS (USANDO GSPREAD) ---
@st.cache_resource
def conectar_gsheets():
    try:
        scope = [
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive"
        ]
        
        # Converte o Secrets em um dicionário Python editável
        creds_dict = dict(st.secrets["connections"]["gsheets"])
        
        # Garante que as quebras de linha da chave privada sejam tratadas corretamente
        if "private_key" in creds_dict:
            creds_dict["private_key"] = creds_dict["private_key"].replace("\\n", "\n")
        
        credentials = Credentials.from_service_account_info(creds_dict, scopes=scope)
        client = gspread.authorize(credentials)
        
        # Abre a planilha pela URL
        spreadsheet_url = creds_dict["spreadsheet"]
        sh = client.open_by_url(spreadsheet_url)
        return sh
    except Exception as e:
        st.error(f"Erro ao conectar ao Google Sheets: {e}")
        return None

# Instância da planilha aberta
sh = conectar_gsheets()

# --- FUNÇÕES AUXILIARES PARA LER E SALVAR ABA A ABA ---
def ler_aba(nome_aba):
    try:
        if sh is None: 
            return pd.DataFrame()
        worksheet = sh.worksheet(nome_aba)
        data = worksheet.get_all_records()
        df = pd.DataFrame(data)
        return df.dropna(how="all")
    except Exception as e:
        return pd.DataFrame()

def salvar_aba(df, nome_aba):
    try:
        if sh is None: 
            return False
        try:
            worksheet = sh.worksheet(nome_aba)
        except gspread.exceptions.WorksheetNotFound:
            worksheet = sh.add_worksheet(title=nome_aba, rows="1000", cols="40")
            
        worksheet.clear()
        
        # Prepara a lista com cabeçalhos e valores para o Google Sheets
        df_clean = df.fillna("")
        dados_lista = [df_clean.columns.values.tolist()] + df_clean.values.tolist()
        worksheet.update(dados_lista)
        return True
    except Exception as e:
        st.error(f"Erro ao salvar na aba {nome_aba}: {e}")
        return False

# --- CARGA E CRUZAMENTO INICIAL DOS DADOS ---
def inicializar_planilha_se_vazia():
    # 1. Conselheiros Iniciais
    df_cons = ler_aba("conselheiros")
    if df_cons.empty:
        dados_cons = pd.DataFrame([
            {"id": 1, "nome": "Francisco Miguel Filho", "cargo": "Conselheiro", "telefone": "", "email": "", "regiao": "São Paulo", "observacoes": ""},
            {"id": 2, "nome": "Vanessa Nassif", "cargo": "Conselheira", "telefone": "", "email": "", "regiao": "São Paulo", "observacoes": ""}
        ])
        salvar_aba(dados_cons, "conselheiros")

    # 2. Cronograma Desmembrado (Word)
    df_crono = ler_aba("cronograma_dados")
    if df_crono.empty and os.path.exists("cronograma.docx"):
        try:
            doc = docx.Document("cronograma.docx")
            if len(doc.tables) > 0:
                rows_list = []
                table = doc.tables[0]
                for row in table.rows[3:]:
                    txts = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                    if not txts or len(txts) < 18: continue
                    distrito = txts[0]
                    if not distrito or distrito.upper().startswith("ZONA") or distrito.upper().startswith("DISTRITOS"): continue
                    
                    def sep(val, q=2):
                        if not val or val == '-' or val == '0-': return ['0'] * q
                        p = [x.strip() for x in re.split(r'[-–—]', str(val)) if x.strip()]
                        while len(p) < q: p.append('0')
                        return p[:q]
                    
                    cras, creas = sep(txts[7], 2)
                    ubs, ama = sep(txts[13], 2)
                    ursi, upa = sep(txts[16], 2)
                    cdi, pai = sep(txts[17], 2)
                    ceu, cdc, ccint = sep(txts[18], 3)
                    col18 = txts[19] if len(txts) > 19 else ""
                    ilpi_2, cdia, nci_2 = sep(col18, 3)
                    col19 = txts[20] if len(txts) > 20 else ""

                    rows_list.append({
                        "distrito": distrito, "ano_criacao": txts[1], "no_mapa": txts[2], "pop_total": txts[3],
                        "pop_masc": txts[4], "pop_fem": txts[5], "subpref_sn": txts[6], "cras": cras, "creas": creas,
                        "caei": txts[8].replace('-',''), "nci": txts[9].replace('-',''), "ilpi": txts[10].replace('-',''),
                        "bpc": txts[11], "rank_vun": txts[12], "ubs": ubs, "ama": ama, "idrpg": txts[14], "emad": txts[15],
                        "ursi": ursi, "upa": upa, "cdi": cdi, "pai": pai, "ceu": ceu, "cdc": cdc, "ccint": ccint,
                        "ilpi_2setor": ilpi_2, "cdia": cdia, "nci_2setor": nci_2, "outros_projetos": col19
                    })
                if rows_list:
                    salvar_aba(pd.DataFrame(rows_list), "cronograma_dados")
        except Exception:
            pass

    # 3. Subprefeituras Totais (Excel)
    df_sub = ler_aba("subprefeituras_dados")
    if df_sub.empty and os.path.exists("Tabela Geral de Registros 2024 RECUPERADO (1).xlsx"):
        try:
            df_t = pd.read_excel("Tabela Geral de Registros 2024 RECUPERADO (1).xlsx", sheet_name="TOTAIS").dropna(how="all")
            if 'Unnamed: 0' in df_t.columns: df_t = df_t.drop(columns=['Unnamed: 0'])
            salvar_aba(df_t, "subprefeituras_dados")
        except Exception:
            pass

    # 4. Registros Base (Excel)
    df_reg = ler_aba("registros_base")
    if df_reg.empty and os.path.exists("Tabela Geral de Registros 2024 RECUPERADO (1).xlsx"):
        try:
            df_b = pd.read_excel("Tabela Geral de Registros 2024 RECUPERADO (1).xlsx", sheet_name="Base de Dados", header=1).dropna(how="all")
            salvar_aba(df_b, "registros_base")
        except Exception:
            pass

# Executa a carga inicial automática se as abas estiverem vazias
inicializar_planilha_se_vazia()

# --- SENHAS DE ACESSO AO MODO EDIÇÃO ---
SENHAS_VALIDAS = ["kico21688", "res1aaa", "res2aaa"]

# Controle de estado do login
if "modo_edicao" not in st.session_state:
    st.session_state["modo_edicao"] = False

# --- FUNÇÃO AUXILIAR PARA GERAR RELATÓRIOS EM EXCEL (.XLSX) ---
def df_para_excel(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Relatorio')
    return output.getvalue()

# Mapeamento de meses para ordenação numérica
MESES_MAPA = {
    "Janeiro": 1, "Fevereiro": 2, "Março": 3, "Abril": 4,
    "Maio": 5, "Junho": 6, "Julho": 7, "Agosto": 8,
    "Setembro": 9, "Outubro": 10, "Novembro": 11, "Dezembro": 12
}

# --- BARRA LATERAL (MODO EDIÇÃO COM SENHA) ---
with st.sidebar:
    st.markdown("### 🔒 Controle de Acesso")
    
    if not st.session_state["modo_edicao"]:
        st.info("👁️ **Modo Consulta (Leitura)**")
        senha_input = st.text_input("Digite a senha para editar:", type="password", key="input_senha")
        if st.button("🔓 Ativar Modo Edição"):
            if senha_input in SENHAS_VALIDAS:
                st.session_state["modo_edicao"] = True
                st.success("Modo Edição Ativado!")
                st.rerun()
            else:
                st.error("Senha incorreta!")
    else:
        st.success("✏️ **Modo Edição Ativo**")
        if st.button("🔒 Sair do Modo Edição"):
            st.session_state["modo_edicao"] = False
            st.rerun()

# --- CABEÇALHO ---
if os.path.exists("logo.png"):
    st.image("logo.png", use_container_width=True)

st.subheader("Painel Geral de Gestão: Políticas e Atenção ao Idoso - SP")
st.markdown("---")

# --- ABAS DA APLICAÇÃO ---
aba_crono, aba_subpref, aba_conselheiros, aba_anotacoes, aba_registros, aba_mapa, aba_noticias, aba_sobre = st.tabs([
    "📋 Cronograma (Distritos)",
    "🏛️ Subprefeituras",
    "👥 Conselheiros Municipais",
    "📝 Anotações Importantes",
    "📌 Registros / Casas de Repouso",
    "🗺️ Fotos, Mapas e Legendas",
    "📰 Notícias e Publicações",
    "ℹ️ Sobre"
])

# -------------------------------------------------------------
# ABA 1: CRONOGRAMA
# -------------------------------------------------------------
with aba_crono:
    st.subheader("Cronograma por Distrito")
    
    df_crono_db = ler_aba("cronograma_dados")

    if st.session_state["modo_edicao"]:
        with st.expander("🛠️ Gerenciar Colunas (Adicionar / Apagar)"):
            c_add, c_del = st.columns(2)
            with c_add:
                st.markdown("**Adicionar Nova Coluna:**")
                nova_col = st.text_input("Nome da nova coluna:", key="crono_add_col")
                if st.button("➕ Adicionar Coluna no Cronograma"):
                    if nova_col and nova_col not in df_crono_db.columns:
                        df_crono_db[nova_col] = ""
                        salvar_aba(df_crono_db, "cronograma_dados")
                        st.success(f"Coluna '{nova_col}' adicionada com sucesso!")
                        st.rerun()
                    elif nova_col in df_crono_db.columns:
                        st.warning("Esta coluna já existe.")
            
            with c_del:
                st.markdown("**Apagar Coluna Existente:**")
                cols_opcoes = [c for c in df_crono_db.columns if c not in ['id', 'ID']]
                col_para_apagar = st.selectbox("Selecione a coluna para remover:", cols_opcoes, key="crono_del_col")
                if st.button("🗑️ Apagar Coluna Selecionada", key="btn_del_crono"):
                    if col_para_apagar:
                        df_crono_db = df_crono_db.drop(columns=[col_para_apagar])
                        salvar_aba(df_crono_db, "cronograma_dados")
                        st.success(f"Coluna '{col_para_apagar}' removida com sucesso!")
                        st.rerun()

    col_filtro, _ = st.columns([1, 2])
    with col_filtro:
        distritos_lista = ["Todos os Distritos"] + sorted(list(df_crono_db['distrito'].dropna().astype(str).unique())) if 'distrito' in df_crono_db.columns else ["Todos os Distritos"]
        distrito_selecionado = st.selectbox("🔍 Selecione o Distrito para Filtrar:", distritos_lista)

    if not st.session_state["modo_edicao"]:
        st.info("ℹ️ Tabela em modo de leitura. Para editar valores, insira a senha na barra lateral.")

    df_exibicao = df_crono_db[df_crono_db['distrito'].astype(str) == distrito_selecionado] if (distrito_selecionado != "Todos os Distritos" and 'distrito' in df_crono_db.columns) else df_crono_db
    
    if st.session_state["modo_edicao"]:
        edited_crono = st.data_editor(df_exibicao, num_rows="dynamic", use_container_width=True, key="editor_crono")
        col_btn1, col_btn2 = st.columns([1, 2])
        with col_btn1:
            if st.button("💾 Salvar Alterações no Cronograma"):
                if distrito_selecionado != "Todos os Distritos" and 'distrito' in df_crono_db.columns:
                    df_crono_db = df_crono_db[df_crono_db['distrito'].astype(str) != distrito_selecionado]
                    df_final = pd.concat([df_crono_db, edited_crono], ignore_index=True)
                else:
                    df_final = edited_crono
                salvar_aba(df_final, "cronograma_dados")
                st.success("Alterações salvas com sucesso no Google Sheets!")
                st.rerun()
        with col_btn2:
            excel_crono = df_para_excel(edited_crono)
            st.download_button("🖨️ Baixar Relatório Filtrado (Excel)", excel_crono, f"Cronograma_{distrito_selecionado}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    else:
        st.dataframe(df_exibicao, use_container_width=True)
        excel_crono = df_para_excel(df_exibicao)
        st.download_button("🖨️ Baixar Relatório Filtrado (Excel)", excel_crono, f"Cronograma_{distrito_selecionado}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# -------------------------------------------------------------
# ABA 2: SUBPREFEITURAS
# -------------------------------------------------------------
with aba_subpref:
    st.subheader("Totais e Indicadores das Subprefeituras")
    
    df_sub_db = ler_aba("subprefeituras_dados")

    if st.session_state["modo_edicao"]:
        with st.expander("🛠️ Gerenciar Colunas (Adicionar / Apagar)"):
            s_add, s_del = st.columns(2)
            with s_add:
                st.markdown("**Adicionar Nova Coluna:**")
                nova_col_sub = st.text_input("Nome da nova coluna:", key="sub_add_col")
                if st.button("➕ Adicionar Coluna em Subprefeituras"):
                    if nova_col_sub and nova_col_sub not in df_sub_db.columns:
                        df_sub_db[nova_col_sub] = ""
                        salvar_aba(df_sub_db, "subprefeituras_dados")
                        st.success(f"Coluna '{nova_col_sub}' adicionada com sucesso!")
                        st.rerun()
                    elif nova_col_sub in df_sub_db.columns:
                        st.warning("Esta coluna já existe.")
            
            with s_del:
                st.markdown("**Apagar Coluna Existente:**")
                cols_sub_opcoes = [c for c in df_sub_db.columns if c not in ['id', 'ID']]
                col_del_sub = st.selectbox("Selecione a coluna para remover:", cols_sub_opcoes, key="sub_del_col")
                if st.button("🗑️ Apagar Coluna Selecionada", key="btn_del_sub"):
                    if col_del_sub:
                        df_sub_db = df_sub_db.drop(columns=[col_del_sub])
                        salvar_aba(df_sub_db, "subprefeituras_dados")
                        st.success(f"Coluna '{col_del_sub}' removida com sucesso!")
                        st.rerun()

    if not st.session_state["modo_edicao"]:
        st.info("ℹ️ Tabela em modo de leitura. Para editar valores, insira a senha na barra lateral.")
        st.dataframe(df_sub_db, use_container_width=True)
        excel_sub = df_para_excel(df_sub_db)
        st.download_button("🖨️ Baixar Relatório de Subprefeituras (Excel)", excel_sub, "Subprefeituras_Totais.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    else:
        edited_sub = st.data_editor(df_sub_db, num_rows="dynamic", use_container_width=True, key="editor_sub")
        col_s1, col_s2 = st.columns([1, 2])
        with col_s1:
            if st.button("💾 Salvar Alterações das Subprefeituras"):
                salvar_aba(edited_sub, "subprefeituras_dados")
                st.success("Dados das Subprefeituras atualizados com sucesso!")
                st.rerun()
        with col_s2:
            excel_sub = df_para_excel(edited_sub)
            st.download_button("🖨️ Baixar Relatório de Subprefeituras (Excel)", excel_sub, "Subprefeituras_Totais.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# -------------------------------------------------------------
# ABA 3: CONSELHEIROS MUNICIPAIS
# -------------------------------------------------------------
with aba_conselheiros:
    st.subheader("Conselheiros Municipais do Conselho do Idoso")
    
    df_cons_db = ler_aba("conselheiros")

    if st.session_state["modo_edicao"]:
        col_cad, col_list = st.columns([1, 2])
        with col_cad:
            st.markdown("### Cadastrar Novo Conselheiro")
            with st.form("form_conselheiro", clear_on_submit=True):
                nome = st.text_input("Nome Completo:")
                cargo = st.text_input("Cargo / Função:", value="Conselheiro(a)")
                telefone = st.text_input("Telefone / WhatsApp:")
                email = st.text_input("E-mail de Contato:")
                regiao = st.text_input("Região / Subprefeitura:")
                obs = st.text_area("Observações:")
                
                submitted = st.form_submit_button("➕ Cadastrar Conselheiro")
                if submitted and nome:
                    novo_id = len(df_cons_db) + 1 if not df_cons_db.empty else 1
                    novo_cons = pd.DataFrame([{
                        "id": novo_id, "nome": nome, "cargo": cargo,
                        "telefone": telefone, "email": email,
                        "regiao": regiao, "observacoes": obs
                    }])
                    df_final_cons = pd.concat([df_cons_db, novo_cons], ignore_index=True)
                    salvar_aba(df_final_cons, "conselheiros")
                    st.success(f"Conselheiro {nome} cadastrado com sucesso!")
                    st.rerun()
    else:
        st.info("ℹ️ Para cadastrar ou excluir conselheiros, ative o Modo Edição na barra lateral.")
        col_list = st.container()

    with col_list:
        st.markdown("### Lista de Conselheiros Cadastrados")
        if not df_cons_db.empty:
            for idx, r in df_cons_db.iterrows():
                c_id = r.get('id', idx)
                c_nome = r.get('nome', '')
                c_cargo = r.get('cargo', '')
                c_tel = r.get('telefone', '')
                c_email = r.get('email', '')
                c_regiao = r.get('regiao', '')
                c_obs = r.get('observacoes', '')
                
                with st.expander(f"👤 {c_nome} - {c_cargo} ({c_regiao or 'SP'})"):
                    st.write(f"**Telefone:** {c_tel or 'Não informado'}")
                    st.write(f"**E-mail:** {c_email or 'Não informado'}")
                    st.write(f"**Observações:** {c_obs or '-'}")
                    
                    if st.session_state["modo_edicao"]:
                        if st.button("❌ Excluir Conselheiro", key=f"del_cons_{idx}"):
                            df_cons_db = df_cons_db.drop(index=idx)
                            salvar_aba(df_cons_db, "conselheiros")
                            st.success("Conselheiro removido com sucesso!")
                            st.rerun()
        else:
            st.info("Nenhum conselheiro cadastrado.")

# -------------------------------------------------------------
# ABA 4: ANOTAÇÕES IMPORTANTES
# -------------------------------------------------------------
with aba_anotacoes:
    st.subheader("Bloco de Anotações e Lembretes Importantes")
    
    df_anot_db = ler_aba("anotacoes")

    if st.session_state["modo_edicao"]:
        col_anot_form, col_anot_view = st.columns([1, 2])
        with col_anot_form:
            st.markdown("### Nova Anotação")
            with st.form("form_anotacao", clear_on_submit=True):
                titulo = st.text_input("Título da Anotação:")
                categoria = st.selectbox("Categoria:", ["Geral", "Reunião", "Vistoria ILPI", "Atendimento", "Outros"])
                conteudo = st.text_area("Conteúdo da Anotação / Lembrete:", height=150)
                
                salvar_anot = st.form_submit_button("📌 Salvar Anotação")
                if salvar_anot and titulo:
                    novo_id = len(df_anot_db) + 1 if not df_anot_db.empty else 1
                    data_hoje = pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
                    nova_anot = pd.DataFrame([{
                        "id": novo_id, "titulo": titulo, "categoria": categoria,
                        "conteudo": conteudo, "data_criacao": data_hoje
                    }])
                    df_final_anot = pd.concat([df_anot_db, nova_anot], ignore_index=True)
                    salvar_aba(df_final_anot, "anotacoes")
                    st.success("Anotação salva com sucesso!")
                    st.rerun()
    else:
        st.info("ℹ️ Para incluir novas anotações ou apagar existentes, ative o Modo Edição na barra lateral.")
        col_anot_view = st.container()

    with col_anot_view:
        st.markdown("### Anotações Salvas")
        if not df_anot_db.empty:
            for idx, r in df_anot_db.iterrows():
                a_tit = r.get('titulo', '')
                a_cat = r.get('categoria', '')
                a_cont = r.get('conteudo', '')
                a_data = str(r.get('data_criacao', ''))
                
                with st.expander(f"📝 [{a_cat}] {a_tit} - {a_data[:16]}"):
                    st.write(a_cont)
                    if st.session_state["modo_edicao"]:
                        if st.button("🗑️ Excluir Anotação", key=f"del_anot_{idx}"):
                            df_anot_db = df_anot_db.drop(index=idx)
                            salvar_aba(df_anot_db, "anotacoes")
                            st.success("Anotação removida com sucesso!")
                            st.rerun()
        else:
            st.info("Nenhuma anotação cadastrada.")

# -------------------------------------------------------------
# ABA 5: REGISTROS / CASAS DE REPOUSO
# -------------------------------------------------------------
with aba_registros:
    st.subheader("Base de Registros e Equipamentos")
    
    df_reg_db = ler_aba("registros_base")

    if st.session_state["modo_edicao"]:
        with st.expander("🛠️ Gerenciar Colunas (Adicionar / Apagar)"):
            r_add, r_del = st.columns(2)
            with r_add:
                st.markdown("**Adicionar Nova Coluna:**")
                nova_col_reg = st.text_input("Nome da nova coluna:", key="reg_add_col")
                if st.button("➕ Adicionar Coluna nos Registros"):
                    if nova_col_reg and nova_col_reg not in df_reg_db.columns:
                        df_reg_db[nova_col_reg] = ""
                        salvar_aba(df_reg_db, "registros_base")
                        st.success(f"Coluna '{nova_col_reg}' adicionada com sucesso!")
                        st.rerun()
                    elif nova_col_reg in df_reg_db.columns:
                        st.warning("Esta coluna já existe.")
            
            with r_del:
                st.markdown("**Apagar Coluna Existente:**")
                cols_reg_opcoes = [c for c in df_reg_db.columns if c not in ['id', 'ID']]
                col_del_reg = st.selectbox("Selecione a coluna para remover:", cols_reg_opcoes, key="reg_del_col")
                if st.button("🗑️ Apagar Coluna Selecionada", key="btn_del_reg"):
                    if col_del_reg:
                        df_reg_db = df_reg_db.drop(columns=[col_del_reg])
                        salvar_aba(df_reg_db, "registros_base")
                        st.success(f"Coluna '{col_del_reg}' removida com sucesso!")
                        st.rerun()

    c1, c2 = st.columns(2)
    with c1:
        subpref_opts = ["Todas"] + sorted([str(x) for x in df_reg_db['subprefeitura'].dropna().unique() if str(x) != 'nan']) if 'subprefeitura' in df_reg_db.columns else ["Todas"]
        subpref_sel = st.selectbox("Subprefeitura:", subpref_opts)
    with c2:
        if 'bairro' in df_reg_db.columns:
            bairros_disp = df_reg_db['bairro'].dropna().unique() if subpref_sel == "Todas" else df_reg_db[df_reg_db['subprefeitura'] == subpref_sel]['bairro'].dropna().unique()
            bairro_opts = ["Todos"] + sorted([str(x) for x in bairros_disp if str(x) != 'nan'])
        else:
            bairro_opts = ["Todos"]
        bairro_sel = st.selectbox("Bairro:", bairro_opts)
        
    df_reg_filt = df_reg_db.copy()
    if subpref_sel != "Todas" and 'subprefeitura' in df_reg_filt.columns:
        df_reg_filt = df_reg_filt[df_reg_filt['subprefeitura'] == subpref_sel]
    if bairro_sel != "Todos" and 'bairro' in df_reg_filt.columns:
        df_reg_filt = df_reg_filt[df_reg_filt['bairro'] == bairro_sel]
        
    if not st.session_state["modo_edicao"]:
        st.info("ℹ️ Tabela em modo de leitura. Para editar valores, insira a senha na barra lateral.")
        st.dataframe(df_reg_filt, use_container_width=True)
        excel_reg = df_para_excel(df_reg_filt)
        st.download_button("🖨️ Baixar Registros Filtrados (Excel)", excel_reg, f"Registros_{subpref_sel}_{bairro_sel}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    else:
        edited_reg = st.data_editor(df_reg_filt, num_rows="dynamic", use_container_width=True, key="editor_reg")
        col_r1, col_r2 = st.columns([1, 2])
        with col_r1:
            if st.button("💾 Salvar Alterações nos Registros"):
                salvar_aba(edited_reg, "registros_base")
                st.success("Registros atualizados com sucesso!")
                st.rerun()
        with col_r2:
            excel_reg = df_para_excel(edited_reg)
            st.download_button("🖨️ Baixar Registros Filtrados (Excel)", excel_reg, f"Registros_{subpref_sel}_{bairro_sel}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# -------------------------------------------------------------
# ABA 6: MAPA E LEGENDAS
# -------------------------------------------------------------
with aba_mapa:
    st.subheader("🗺️ Fotos, Mapas e Legendas Oficiais")
    
    @st.dialog("🔍 Visualização Ampliada da Imagem", width="large")
    def popup_imagem(caminho_ou_bytes, titulo_img="Imagem"):
        st.markdown(f"### {titulo_img}")
        zoom_px = st.slider("🔍 Ajuste o Tamanho / Zoom da Imagem (Pixels):", min_value=600, max_value=2200, value=1100, step=100, key="modal_zoom")
        st.image(caminho_ou_bytes, width=zoom_px)
        st.markdown("---")
        if st.button("❌ Fechar Imagem", use_container_width=True):
            st.rerun()

    col_mapa_main, col_upload = st.columns([1, 1])
    
    with col_mapa_main:
        st.markdown("### Mapa de São Paulo e Legenda Padrão")
        nome_imagem = "mapa_legendas.png"
        if os.path.exists(nome_imagem):
            st.image(nome_imagem, width=280, caption="Clique no botão abaixo para expandir em tela cheia")
            if st.button("🔎 Ampliar Mapa em Tela Cheia", key="btn_open_mapa"):
                popup_imagem(nome_imagem, "Mapa de SP e Legendas")
        else:
            st.warning(f"O arquivo '{nome_imagem}' não foi localizado na pasta do aplicativo.")

    with col_upload:
        st.info("ℹ️ Para adicionar novas fotos/mapas permanentes, salve as imagens diretamente no repositório do projeto.")

# -------------------------------------------------------------
# ABA 7: NOTÍCIAS E PUBLICAÇÕES (PDF / DOC)
# -------------------------------------------------------------
with aba_noticias:
    st.subheader("📰 Notícias, Informativos e Publicações")

    df_noticias = ler_aba("noticias_pdf")

    @st.dialog("🔍 Leitor e Visualizador de Documento", width="large")
    def popup_pdf(bytes_arquivo, nome_arq, titulo_doc):
        st.markdown(f"### {titulo_doc}")
        if nome_arq.lower().endswith('.pdf'):
            st.info("Abaixo estão as páginas do documento convertidas em imagem para facilitar a sua leitura:")
            try:
                doc = fitz.open(stream=bytes_arquivo, filetype="pdf")
                for i in range(len(doc)):
                    page = doc.load_page(i)
                    pix = page.get_pixmap(dpi=150)
                    img_bytes = pix.tobytes("png")
                    st.image(img_bytes, caption=f"Página {i + 1} de {len(doc)}", use_container_width=True)
                    st.markdown("---")
            except Exception as e:
                st.error("Não foi possível processar a imagem deste documento PDF.")
            
            st.download_button("📥 Baixar o PDF Original", bytes_arquivo, nome_arq, use_container_width=True)
        else:
            st.info("Visualização em tela cheia não suportada para arquivos do Word. Utilize o botão abaixo para baixá-lo:")
            st.download_button("📥 Baixar Documento (Word)", bytes_arquivo, nome_arq, use_container_width=True)
            
        if st.button("❌ Fechar Visualizador", use_container_width=True):
            st.rerun()

    col_not_view, col_not_up = st.columns([2, 1])

    with col_not_up:
        if st.session_state["modo_edicao"]:
            st.markdown("### 📥 Publicar Nova Notícia / Documento")
            with st.form("form_noticia", clear_on_submit=True):
                tit_noticia = st.text_input("Título da Publicação / Informativo:")
                
                col_m, col_a = st.columns(2)
                with col_m:
                    mes_sel = st.selectbox("Mês de Referência:", list(MESES_MAPA.keys()))
                with col_a:
                    ano_sel = st.number_input("Ano de Referência:", min_value=2020, max_value=2035, value=2026, step=1)
                
                arq_up = st.file_uploader("Selecione o arquivo (PDF ou Word):", type=["pdf", "docx", "doc"])
                pub_btn = st.form_submit_button("📤 Publicar Notícia")

                if pub_btn and arq_up and tit_noticia:
                    m_num = MESES_MAPA[mes_sel]
                    
                    # Salva o arquivo localmente em pasta de upload para persistência
                    os.makedirs("uploads_noticias", exist_ok=True)
                    caminho_salvo = os.path.join("uploads_noticias", arq_up.name)
                    with open(caminho_salvo, "wb") as f:
                        f.write(arq_up.read())
                    
                    novo_id = len(df_noticias) + 1 if not df_noticias.empty else 1
                    nova_noticia = pd.DataFrame([{
                        "id": novo_id, "titulo": tit_noticia, "mes_ref": mes_sel,
                        "ano_ref": ano_sel, "mes_num": m_num, "nome_arquivo": arq_up.name,
                        "data_upload": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
                    }])
                    
                    df_final_not = pd.concat([df_noticias, nova_noticia], ignore_index=True)
                    salvar_aba(df_final_not, "noticias_pdf")
                    st.success("Publicação cadastrada com sucesso!")
                    st.rerun()
        else:
            st.info("ℹ️ Para cadastrar novas publicações ou notícias, ative o Modo Edição na barra lateral.")

    with col_not_view:
        st.markdown("### 📚 Publicações Recentes (Ordenadas por Mês/Ano)")
        
        if not df_noticias.empty:
            df_noticias_ord = df_noticias.sort_values(by=["ano_ref", "mes_num"], ascending=[False, False])
            
            for idx, r in df_noticias_ord.iterrows():
                n_tit = r.get('titulo', '')
                n_mes = r.get('mes_ref', '')
                n_ano = r.get('ano_ref', '')
                n_nome_arq = r.get('nome_arquivo', '')
                
                with st.expander(f"📄 [{n_mes} / {n_ano}] - {n_tit}"):
                    st.write(f"**Arquivo:** `{n_nome_arq}`")
                    
                    caminho_local = os.path.join("uploads_noticias", str(n_nome_arq))
                    if os.path.exists(caminho_local):
                        with open(caminho_local, "rb") as f_bytes:
                            bytes_arq = f_bytes.read()
                        
                        b_ler, b_down, b_del = st.columns([1, 1, 1])
                        with b_ler:
                            if st.button("🔍 Abrir / Ler Janela", key=f"read_pdf_{idx}"):
                                popup_pdf(bytes_arq, str(n_nome_arq), n_tit)
                        with b_down:
                            mime_t = "application/pdf" if str(n_nome_arq).lower().endswith('.pdf') else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            st.download_button("📥 Baixar File", bytes_arq, file_name=str(n_nome_arq), mime=mime_t, key=f"down_pdf_{idx}")
                        with b_del:
                            if st.session_state["modo_edicao"]:
                                if st.button("🗑️ Apagar Publicação", key=f"del_pdf_{idx}"):
                                    df_noticias = df_noticias.drop(index=idx)
                                    salvar_aba(df_noticias, "noticias_pdf")
                                    st.success("Notícia removida com sucesso!")
                                    st.rerun()
                    else:
                        st.warning("O arquivo original desta notícia não foi localizado no servidor.")
                        if st.session_state["modo_edicao"]:
                            if st.button("🗑️ Remover Registro", key=f"del_rec_{idx}"):
                                df_noticias = df_noticias.drop(index=idx)
                                salvar_aba(df_noticias, "noticias_pdf")
                                st.rerun()
        else:
            st.info("Nenhuma publicação ou notícia enviada ainda. Ative o Modo Edição para subir novos PDFs.")

# -------------------------------------------------------------
# ABA 8: SOBRE (ÚLTIMA ABA)
# -------------------------------------------------------------
with aba_sobre:
    st.subheader("ℹ️ Sobre esta Aplicação")
    st.markdown('''
    Este painel foi desenvolvido especialmente para apoio à consulta, gestão e acompanhamento de políticas públicas voltadas às pessoas idosas na cidade de São Paulo.
    
    * **👨‍💻 Criadores:** Rodrigo Prado Miguel & Francisco Miguel Filho
    * **⚙️ Desenvolvimento e Organização:** Aplicação integrada com armazenamento seguro no Google Sheets.
    * **📊 Bases Utilizadas:** Tabela Geral de Registros 2024 e Cronograma de Indicadores por Distrito.
    * **🏷️ Versão:** 2.0.0 (Google Sheets Storage)
    ''')
